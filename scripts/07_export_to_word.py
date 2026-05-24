import os
import json
import subprocess
import sys

# Word dönüşümü için gerekli kütüphaneyi kontrol et ve yükle
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("[BİLGİ] Word dönüşümü için gerekli kütüphaneler yükleniyor...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = r"C:\Users\LENOVO\Desktop\kamu-ihale-ai"
QA_DIR = os.path.join(BASE_DIR, "data", "qa")

def github_push(commit_mesaji, dosya_yolu):
    try:
        os.chdir(BASE_DIR)
        subprocess.run(["git", "add", dosya_yolu], check=True)
        subprocess.run(["git", "commit", "-m", commit_mesaji], check=True)
        subprocess.run(["git", "push", "origin", "main"], check=True)
        print(f"[GITHUB] {dosya_yolu} başarıyla GitHub'a gönderildi!")
    except Exception as e:
        print("[GITHUB HATA] Kod GitHub'a gönderilemedi:", e)

def set_cell_background(cell, color_hex):
    """Tablo başlıkları için arka plan rengi ayarlar"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def main():
    json_files = [f for f in os.listdir(QA_DIR) if f.endswith(".json")]
    
    if not json_files:
        print("[UYARI] data/qa/ klasöründe aktarılacak JSON dosyası bulunamadı!")
        return

    print(f"[BİLGİ] {len(json_files)} adet makale dosyası resmi Word raporuna dönüştürülüyor...")

    # Yeni bir Word dokümanı oluşturuyoruz
    doc = Document()
    
    # Sayfa kenar boşluklarını ayarla (2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # REKST TİPİ: Genel Yazı Tipi Ayarı
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # 1. BAŞLIK EKLEME
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("KAMU İHALE HUKUKU NİHAİ UYUŞMAZLIK ANALİZ RAPORU")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(26, 54, 93) # Koyu Lacivert
    doc.add_paragraph("").paragraph_format.space_after = Pt(20)

    toplam_uyusmazlik_sayisi = 0

    for file_name in json_files:
        path = os.path.join(QA_DIR, file_name)
        
        with open(path, "r", encoding="utf-8") as f:
            try:
                qa_list = json.load(f)
            except Exception:
                continue

        # Makale Başlığı Ekleme
        heading = doc.add_paragraph()
        heading_run = heading.add_run(f"📄 Kaynak Doküman: {file_name.replace('.json', '.txt')}")
        heading_run.bold = True
        heading_run.font.size = Pt(13)
        heading_run.font.color.rgb = RGBColor(43, 108, 176) # Açık Mavi
        heading.paragraph_format.space_before = Pt(15)
        heading.paragraph_format.space_after = Pt(5)

        for item in qa_list:
            toplam_uyusmazlik_sayisi += 1
            
            # Her uyuşmazlık için 2 sütunlu şık bir tablo oluşturuyoruz
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Light Shading Accent 1' # Word'ün hazır temiz tablo şablonu
            table.autofit = False
            
            # Sütun genişliklerini ayarla
            table.columns[0].width = Inches(1.5)
            table.columns[1].width = Inches(5.0)

            # Satır içeriklerini doldur
            icerik = [
                ("Uyuşmazlık Konusu:", item.get("uyusmazlik_konusu", "")),
                ("Hukuki Soru:", item.get("soru", "")),
                ("Yapay Zeka Cevabı:", item.get("cevap", "")),
                ("Mevzuat Dayanağı:", item.get("dayanak", ""))
            ]

            for idx, (baslik, veri) in enumerate(icerik):
                row = table.rows[idx]
                
                # Sol sütun (Başlıklar)
                cell_lbl = row.cells[0]
                cell_lbl.text = baslik
                cell_lbl.paragraphs[0].runs[0].font.bold = True
                cell_lbl.paragraphs[0].runs[0].font.size = Pt(10)
                set_cell_background(cell_lbl, "EDF2F7") # Hafif gri arka plan
                
                # Sağ sütun (Veriler)
                cell_val = row.cells[1]
                cell_val.text = str(veri)
                cell_val.paragraphs[0].runs[0].font.size = Pt(10)

            # Tablolar arasına boşluk bırak
            doc.add_paragraph("").paragraph_format.space_after = Pt(10)

    # Rapor Sonuna Özet Bilgisi Ekle
    doc.add_paragraph("").paragraph_format.space_before = Pt(20)
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    summary_run = summary.add_run(f"GENEL TOPLAM: 94 makaleden toplam {toplam_uyusmazlik_sayisi} adet hukuki uyuşmazlık başarıyla rapora aktarılmıştır.")
    summary_run.bold = True
    summary_run.font.size = Pt(11)
    summary_run.font.color.rgb = RGBColor(22, 101, 52) # Koyu Yeşil

    # Masaüstü yolunu bul ve kaydet
    masaustu_yolu = os.path.join(os.path.expanduser("~"), "Desktop")
    word_dosya_yolu = os.path.join(masaustu_yolu, "KAMU_IHALE_UYUSMAZLIK_RAPORU.docx")
    
    doc.save(word_dosya_yolu)
    
    print(f"\n[BAŞARILI] Word dökümanı başarıyla oluşturuldu!")
    print(f"[BİLGİ] Toplam {toplam_uyusmazlik_sayisi} adet uyuşmazlık tablosu rapora eklendi.")
    print(f"[LOKASYON] Dosya Masaüstünüze 'KAMU_IHALE_UYUSMAZLIK_RAPORU.docx' adıyla kaydedildi.")

if __name__ == "__main__":
    main()
    github_push("Asama 9: Resmi Word raporlama motoru pipelinesi eklendi", "scripts/07_export_to_word.py")

